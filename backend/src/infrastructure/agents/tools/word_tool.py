"""
Word Document Tool — generates professional .docx files from structured section data.

Architecture:
  User → Agent → create_word_doc tool
                    ↓
                SectionData (title + content with markdown)
                    ↓
                python-docx Document (with rich formatting + charts)
                    ↓
                storage/word_docs/<uuid>.docx
                    ↓
                Download URL

Supported Markdown in content:
  - ## Sub-Heading / ### Sub-Sub-Heading
  - **bold**, *italic*, ***bold italic***
  - | col1 | col2 | table syntax
  - - bullet / * bullet
  - 1. numbered items
  - > blockquote
  - --- horizontal rule
  - ```chart ... ``` blocks (bar, line, pie)

Resource limits (enforced here, not by the LLM):
  - Max sections   : 50
  - Max file size   : 10 MB
"""

import os
import re
import io
import math
import uuid
import logging
import tempfile
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

from pydantic import BaseModel, Field
from sqlalchemy.orm import Session
from langchain_core.tools import StructuredTool
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

try:
    from PIL import Image, ImageDraw, ImageFont
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    logger_pil = logging.getLogger(__name__)
    logger_pil.warning("Pillow not installed — chart rendering disabled.")

from src.infrastructure.database.models import User, GeneratedWord, WordSection

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Resource limits
# ---------------------------------------------------------------------------

MAX_SECTIONS = 50
MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB

# ---------------------------------------------------------------------------
# Color Palette — Professional dark-blue theme
# ---------------------------------------------------------------------------

COLORS = {
    "primary":       RGBColor(0x1B, 0x2A, 0x4A),   # Dark navy
    "secondary":     RGBColor(0x2C, 0x5F, 0x8A),   # Steel blue
    "accent":        RGBColor(0x3A, 0x86, 0xC8),   # Bright blue
    "heading_1":     RGBColor(0x1B, 0x2A, 0x4A),   # Navy
    "heading_2":     RGBColor(0x2C, 0x5F, 0x8A),   # Steel blue
    "heading_3":     RGBColor(0x3A, 0x86, 0xC8),   # Bright blue
    "chart_bg":      RGBColor(0xFF, 0xFF, 0xFF),   # White background
    "body":          RGBColor(0x2D, 0x2D, 0x2D),   # Dark gray
    "light_gray":    RGBColor(0x6B, 0x70, 0x7B),   # Muted gray
    "table_header":  RGBColor(0x1B, 0x2A, 0x4A),   # Navy
    "table_alt_row": RGBColor(0xF0, 0xF4, 0xF8),   # Light blue-gray
    "border":        RGBColor(0xD0, 0xD5, 0xDD),   # Border gray
    "blockquote":    RGBColor(0x4A, 0x5A, 0x6A),   # Quote gray
    "white":         RGBColor(0xFF, 0xFF, 0xFF),
}

# ---------------------------------------------------------------------------
# Input Schemas
# ---------------------------------------------------------------------------


class WordSectionInput(BaseModel):
    """Represents one section in the Word document."""
    title: str = Field(description="Title/heading of this section.")
    content: str = Field(
        description=(
            "Text content for this section. Supports markdown-like formatting: "
            "**bold**, *italic*, ## Sub-Heading, ### Sub-Sub-Heading, "
            "- bullet points, 1. numbered items, > blockquotes, "
            "--- horizontal rules, and | col1 | col2 | table syntax."
        )
    )
    section_type: str = Field(
        default="text",
        description="Type of section: text, list, table"
    )


class CreateWordInput(BaseModel):
    """Input schema for the create_word_doc tool."""
    title: str = Field(
        description="Title / filename label for the Word document."
    )
    sections: List[WordSectionInput] = Field(
        description=(
            "List of sections for the document. Each section has a 'title' and 'content'. "
            "Content supports markdown formatting for rich output."
        )
    )
    base_url: str = Field(
        description=(
            "Base URL of the backend server (e.g. http://localhost:8000). "
            "Injected automatically — do not ask the user."
        )
    )


# ---------------------------------------------------------------------------
# Document styling helpers
# ---------------------------------------------------------------------------


def _setup_styles(doc: Document) -> None:
    """Configure document-wide styles for a professional look."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = COLORS["body"]
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(6)
    paragraph_format.space_before = Pt(2)
    paragraph_format.line_spacing = 1.15

    # Heading 1
    if "Heading 1" in doc.styles:
        h1 = doc.styles["Heading 1"]
        h1.font.name = "Calibri"
        h1.font.size = Pt(18)
        h1.font.bold = True
        h1.font.color.rgb = COLORS["heading_1"]
        h1.paragraph_format.space_before = Pt(18)
        h1.paragraph_format.space_after = Pt(8)
        h1.paragraph_format.keep_with_next = True

    # Heading 2
    if "Heading 2" in doc.styles:
        h2 = doc.styles["Heading 2"]
        h2.font.name = "Calibri"
        h2.font.size = Pt(14)
        h2.font.bold = True
        h2.font.color.rgb = COLORS["heading_2"]
        h2.paragraph_format.space_before = Pt(14)
        h2.paragraph_format.space_after = Pt(6)
        h2.paragraph_format.keep_with_next = True

    # Heading 3
    if "Heading 3" in doc.styles:
        h3 = doc.styles["Heading 3"]
        h3.font.name = "Calibri"
        h3.font.size = Pt(12)
        h3.font.bold = True
        h3.font.color.rgb = COLORS["heading_3"]
        h3.paragraph_format.space_before = Pt(10)
        h3.paragraph_format.space_after = Pt(4)
        h3.paragraph_format.keep_with_next = True


def _add_header_footer(doc: Document, title: str) -> None:
    """Add professional header and footer to all sections."""
    for section in doc.sections:
        # Header — document title, right-aligned, subtle
        header = section.header
        header.is_linked_to_previous = False
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header_para.add_run(title)
        run.font.size = Pt(8)
        run.font.color.rgb = COLORS["light_gray"]
        run.font.name = "Calibri"
        run.font.italic = True

        # Footer — page number, centered
        footer = section.footer
        footer.is_linked_to_previous = False
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add "Page X" with auto page number field
        run = footer_para.add_run("Page ")
        run.font.size = Pt(8)
        run.font.color.rgb = COLORS["light_gray"]
        run.font.name = "Calibri"

        # Page number field code
        fld_xml = (
            f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE \\* MERGEFORMAT ">'
            f'<w:r><w:rPr><w:sz w:val="16"/><w:color w:val="6B707B"/>'
            f'<w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/></w:rPr>'
            f'<w:t>1</w:t></w:r></w:fldSimple>'
        )
        footer_para._element.append(parse_xml(fld_xml))


def _add_cover_page(doc: Document, title: str) -> None:
    """Add a professional cover/title page."""
    # Multiple blank lines to push title to upper-third
    for _ in range(6):
        doc.add_paragraph("")

    # Accent line above title
    line_para = doc.add_paragraph()
    line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line_para.add_run("━" * 40)
    run.font.color.rgb = COLORS["accent"]
    run.font.size = Pt(14)

    # Document title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(12)
    title_para.paragraph_format.space_after = Pt(12)
    run = title_para.add_run(title)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = COLORS["primary"]
    run.font.name = "Calibri"

    # Accent line below title
    line_para2 = doc.add_paragraph()
    line_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = line_para2.add_run("━" * 40)
    run2.font.color.rgb = COLORS["accent"]
    run2.font.size = Pt(14)

    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.paragraph_format.space_before = Pt(24)
    date_run = date_para.add_run(datetime.now().strftime("%B %d, %Y"))
    date_run.font.size = Pt(12)
    date_run.font.color.rgb = COLORS["light_gray"]
    date_run.font.name = "Calibri"

    # Confidential notice
    conf_para = doc.add_paragraph()
    conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf_para.paragraph_format.space_before = Pt(48)
    conf_run = conf_para.add_run("CONFIDENTIAL")
    conf_run.font.size = Pt(9)
    conf_run.font.color.rgb = COLORS["light_gray"]
    conf_run.font.name = "Calibri"
    conf_run.font.italic = True

    # Page break after cover
    doc.add_page_break()


def _add_table_of_contents(doc: Document, sections: List[WordSectionInput]) -> None:
    """Add a table of contents page."""
    toc_heading = doc.add_heading("Table of Contents", level=1)

    # Add spacing
    doc.add_paragraph("")

    for i, sec in enumerate(sections, 1):
        toc_para = doc.add_paragraph()
        toc_para.paragraph_format.space_after = Pt(4)
        toc_para.paragraph_format.space_before = Pt(2)
        
        # Section number and title
        num_run = toc_para.add_run(f"{i}.  ")
        num_run.font.size = Pt(11)
        num_run.font.color.rgb = COLORS["accent"]
        num_run.font.bold = True
        num_run.font.name = "Calibri"

        title_run = toc_para.add_run(sec.title)
        title_run.font.size = Pt(11)
        title_run.font.color.rgb = COLORS["body"]
        title_run.font.name = "Calibri"

    # Separator after TOC
    doc.add_paragraph("")
    doc.add_page_break()


def _add_section_divider(doc: Document) -> None:
    """Add a subtle horizontal rule between sections."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(8)
    run = para.add_run("─" * 60)
    run.font.size = Pt(8)
    run.font.color.rgb = COLORS["border"]


# ---------------------------------------------------------------------------
# Markdown parsing → Word formatting
# ---------------------------------------------------------------------------


def _parse_inline_markdown(paragraph, text: str) -> None:
    """
    Parse inline markdown (bold, italic, bold-italic) and add formatted runs
    to the paragraph. Handles mixed plain + formatted text.
    """
    # Pattern: ***bold italic***, **bold**, *italic*
    pattern = re.compile(
        r'(\*\*\*(.+?)\*\*\*)'    # bold italic
        r'|(\*\*(.+?)\*\*)'       # bold
        r'|(\*(.+?)\*)'           # italic
    )

    last_end = 0
    for match in pattern.finditer(text):
        # Add plain text before this match
        if match.start() > last_end:
            plain = text[last_end:match.start()]
            run = paragraph.add_run(plain)
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            run.font.color.rgb = COLORS["body"]

        if match.group(2):  # ***bold italic***
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.italic = True
        elif match.group(4):  # **bold**
            run = paragraph.add_run(match.group(4))
            run.bold = True
        elif match.group(6):  # *italic*
            run = paragraph.add_run(match.group(6))
            run.italic = True

        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = COLORS["body"]
        last_end = match.end()

    # Remaining plain text after last match
    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = COLORS["body"]

    # If no matches at all, add as plain text
    if last_end == 0 and text:
        # Already handled by the "remaining" block above
        pass


def _parse_table_block(doc: Document, lines: List[str]) -> None:
    """
    Parse a markdown table block and create a styled Word table.
    Lines should be like:
    | Header1 | Header2 | Header3 |
    |---------|---------|---------|
    | Cell1   | Cell2   | Cell3   |
    """
    # Parse rows
    rows_data = []
    separator_idx = -1

    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped.startswith("|"):
            continue
        # Check if it's a separator row (|---|---|)
        if re.match(r'^\|[\s\-:]+(\|[\s\-:]+)*\|?$', stripped):
            separator_idx = i
            continue
        # Parse cells
        cells = [c.strip() for c in stripped.split("|")]
        # Remove empty first/last elements from split
        cells = [c for c in cells if c or cells.index(c) not in (0, len(cells) - 1)]
        # Filter truly empty boundary cells
        if cells and cells[0] == "":
            cells = cells[1:]
        if cells and cells[-1] == "":
            cells = cells[:-1]
        if cells:
            rows_data.append(cells)

    if not rows_data:
        return

    # Determine column count from widest row
    num_cols = max(len(row) for row in rows_data)

    # Pad shorter rows
    for row in rows_data:
        while len(row) < num_cols:
            row.append("")

    # Create table
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Style the table
    for row_idx, row_data in enumerate(rows_data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = ""  # Clear default
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            if row_idx == 0:
                # Header row styling
                run = para.add_run(cell_text)
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = COLORS["white"]
                run.font.name = "Calibri"
                # Background color for header
                shading = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="1B2A4A" w:val="clear"/>'
                )
                cell._element.get_or_add_tcPr().append(shading)
            else:
                run = para.add_run(cell_text)
                run.font.size = Pt(10)
                run.font.color.rgb = COLORS["body"]
                run.font.name = "Calibri"
                # Alternating row colors
                if row_idx % 2 == 0:
                    shading = parse_xml(
                        f'<w:shd {nsdecls("w")} w:fill="F0F4F8" w:val="clear"/>'
                    )
                    cell._element.get_or_add_tcPr().append(shading)

            # Cell padding
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after = Pt(3)

    # Add spacing after table
    doc.add_paragraph("")


def _add_blockquote(doc: Document, text: str) -> None:
    """Add a styled blockquote paragraph."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.5)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)

    # Add left border via XML
    pPr = para._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="12" w:space="8" w:color="3A86C8"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    run = para.add_run(text)
    run.font.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = COLORS["blockquote"]
    run.font.name = "Calibri"


def _add_horizontal_rule(doc: Document) -> None:
    """Add a horizontal rule / separator."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(10)
    run = para.add_run("─" * 70)
    run.font.size = Pt(6)
    run.font.color.rgb = COLORS["border"]


# ---------------------------------------------------------------------------
# Chart rendering (using Pillow)
# ---------------------------------------------------------------------------

# Chart color palette — vibrant, professional
CHART_COLORS = [
    (0x3A, 0x86, 0xC8),   # Blue
    (0xE8, 0x6C, 0x50),   # Coral-red
    (0x2E, 0xCC, 0x71),   # Green
    (0xF3, 0x9C, 0x12),   # Amber
    (0x9B, 0x59, 0xB6),   # Purple
    (0x1A, 0xBC, 0x9C),   # Teal
    (0xE7, 0x4C, 0x3C),   # Red
    (0x34, 0x98, 0xDB),   # Light blue
    (0xF1, 0xC4, 0x0F),   # Yellow
    (0x2C, 0x3E, 0x50),   # Dark slate
]


def _get_text_size(draw: 'ImageDraw.Draw', text: str, font: 'ImageFont.FreeTypeFont') -> Tuple[int, int]:
    """Get text bounding box size, compatible with older Pillow versions."""
    try:
        bbox = draw.textbbox((0, 0), text, font=font)
        return bbox[2] - bbox[0], bbox[3] - bbox[1]
    except AttributeError:
        return draw.textsize(text, font=font)


def _get_font(size: int, bold: bool = False) -> 'ImageFont.FreeTypeFont':
    """Try to load a good font, fall back to default."""
    font_paths = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        "/usr/share/fonts/truetype/freefont/FreeSans.ttf",
    ]
    for path in font_paths:
        try:
            return ImageFont.truetype(path, size)
        except (OSError, IOError):
            continue
    try:
        return ImageFont.truetype("arial.ttf", size)
    except (OSError, IOError):
        return ImageFont.load_default()


def _render_bar_chart(title: str, labels: List[str], values: List[float], width: int = 800, height: int = 500) -> Image.Image:
    """Render a bar chart as a PIL Image."""
    img = Image.new('RGB', (width, height), (255, 255, 255))
    draw = ImageDraw.Draw(img)

    title_font = _get_font(18, bold=True)
    label_font = _get_font(12)
    value_font = _get_font(11, bold=True)

    # Title
    tw, th = _get_text_size(draw, title, title_font)
    draw.text(((width - tw) // 2, 15), title, fill=(0x1B, 0x2A, 0x4A), font=title_font)

    # Chart area
    margin_left = 80
    margin_right = 40
    margin_top = 60
    margin_bottom = 80
    chart_w = width - margin_left - margin_right
    chart_h = height - margin_top - margin_bottom

    if not values:
        return img

    max_val = max(values) if max(values) > 0 else 1
    n = len(labels)
    bar_width = max(20, min(60, chart_w // (n * 2)))
    spacing = (chart_w - bar_width * n) / (n + 1)

    # Y-axis gridlines
    num_gridlines = 5
    for i in range(num_gridlines + 1):
        y = margin_top + chart_h - (i * chart_h / num_gridlines)
        val = (i * max_val / num_gridlines)
        draw.line([(margin_left, int(y)), (width - margin_right, int(y))], fill=(0xE0, 0xE0, 0xE0), width=1)
        val_str = f"{val:,.0f}" if val == int(val) else f"{val:,.1f}"
        vw, vh = _get_text_size(draw, val_str, label_font)
        draw.text((margin_left - vw - 8, int(y) - vh // 2), val_str, fill=(0x6B, 0x70, 0x7B), font=label_font)

    # Bars
    for i, (label, val) in enumerate(zip(labels, values)):
        x = margin_left + spacing + i * (bar_width + spacing)
        bar_h = (val / max_val) * chart_h if max_val > 0 else 0
        y_top = margin_top + chart_h - bar_h
        y_bottom = margin_top + chart_h
        color = CHART_COLORS[i % len(CHART_COLORS)]

        # Bar with slight rounded effect (draw main rect)
        draw.rectangle([int(x), int(y_top), int(x + bar_width), int(y_bottom)], fill=color)

        # Value on top
        val_str = f"{val:,.0f}" if val == int(val) else f"{val:,.1f}"
        vw, vh = _get_text_size(draw, val_str, value_font)
        draw.text((int(x + bar_width / 2 - vw / 2), int(y_top - vh - 4)), val_str, fill=color, font=value_font)

        # Label below
        lw, lh = _get_text_size(draw, label, label_font)
        label_x = int(x + bar_width / 2 - lw / 2)
        draw.text((label_x, int(y_bottom + 8)), label, fill=(0x2D, 0x2D, 0x2D), font=label_font)

    # Axes
    draw.line([(margin_left, margin_top), (margin_left, margin_top + chart_h)], fill=(0x2D, 0x2D, 0x2D), width=2)
    draw.line([(margin_left, margin_top + chart_h), (width - margin_right, margin_top + chart_h)], fill=(0x2D, 0x2D, 0x2D), width=2)

    return img


def _render_line_chart(title: str, labels: List[str], values: List[float], width: int = 800, height: int = 500) -> Image.Image:
    """Render a line chart as a PIL Image."""
    img = Image.new('RGB', (width, height), (255, 255, 255))
    draw = ImageDraw.Draw(img)

    title_font = _get_font(18, bold=True)
    label_font = _get_font(12)
    value_font = _get_font(10, bold=True)

    # Title
    tw, th = _get_text_size(draw, title, title_font)
    draw.text(((width - tw) // 2, 15), title, fill=(0x1B, 0x2A, 0x4A), font=title_font)

    margin_left = 80
    margin_right = 40
    margin_top = 60
    margin_bottom = 80
    chart_w = width - margin_left - margin_right
    chart_h = height - margin_top - margin_bottom

    if not values or len(values) < 2:
        return img

    max_val = max(values) if max(values) > 0 else 1
    min_val = min(values)
    val_range = max_val - min_val if max_val != min_val else 1
    n = len(labels)

    # Y-axis gridlines
    num_gridlines = 5
    for i in range(num_gridlines + 1):
        y = margin_top + chart_h - (i * chart_h / num_gridlines)
        val = min_val + (i * val_range / num_gridlines)
        draw.line([(margin_left, int(y)), (width - margin_right, int(y))], fill=(0xE0, 0xE0, 0xE0), width=1)
        val_str = f"{val:,.0f}" if val == int(val) else f"{val:,.1f}"
        vw, vh = _get_text_size(draw, val_str, label_font)
        draw.text((margin_left - vw - 8, int(y) - vh // 2), val_str, fill=(0x6B, 0x70, 0x7B), font=label_font)

    # Calculate points
    points = []
    step = chart_w / (n - 1) if n > 1 else chart_w
    for i, val in enumerate(values):
        x = margin_left + i * step
        y = margin_top + chart_h - ((val - min_val) / val_range) * chart_h
        points.append((int(x), int(y)))

    # Area fill (light)
    color = CHART_COLORS[0]
    fill_color = (color[0], color[1], color[2], 30)
    area_points = points.copy()
    area_points.append((points[-1][0], margin_top + chart_h))
    area_points.append((points[0][0], margin_top + chart_h))

    # Draw area fill with semi-transparent overlay
    overlay = Image.new('RGBA', (width, height), (255, 255, 255, 0))
    overlay_draw = ImageDraw.Draw(overlay)
    overlay_draw.polygon(area_points, fill=(color[0], color[1], color[2], 40))
    img.paste(Image.alpha_composite(Image.new('RGBA', img.size, (255, 255, 255, 255)), overlay).convert('RGB'))
    draw = ImageDraw.Draw(img)  # Re-create draw after paste

    # Re-draw gridlines on top of fill
    for i in range(num_gridlines + 1):
        y = margin_top + chart_h - (i * chart_h / num_gridlines)
        draw.line([(margin_left, int(y)), (width - margin_right, int(y))], fill=(0xE0, 0xE0, 0xE0), width=1)

    # Line
    for i in range(len(points) - 1):
        draw.line([points[i], points[i + 1]], fill=color, width=3)

    # Data points + values
    for i, (px, py) in enumerate(points):
        draw.ellipse([px - 5, py - 5, px + 5, py + 5], fill=color, outline=(255, 255, 255))
        val_str = f"{values[i]:,.0f}" if values[i] == int(values[i]) else f"{values[i]:,.1f}"
        vw, vh = _get_text_size(draw, val_str, value_font)
        draw.text((px - vw // 2, py - vh - 8), val_str, fill=color, font=value_font)

    # X-axis labels
    for i, label in enumerate(labels):
        x = margin_left + i * step
        lw, lh = _get_text_size(draw, label, label_font)
        draw.text((int(x - lw / 2), margin_top + chart_h + 8), label, fill=(0x2D, 0x2D, 0x2D), font=label_font)

    # Axes
    draw.line([(margin_left, margin_top), (margin_left, margin_top + chart_h)], fill=(0x2D, 0x2D, 0x2D), width=2)
    draw.line([(margin_left, margin_top + chart_h), (width - margin_right, margin_top + chart_h)], fill=(0x2D, 0x2D, 0x2D), width=2)

    return img


def _render_pie_chart(title: str, labels: List[str], values: List[float], width: int = 700, height: int = 500) -> Image.Image:
    """Render a pie chart as a PIL Image."""
    img = Image.new('RGB', (width, height), (255, 255, 255))
    draw = ImageDraw.Draw(img)

    title_font = _get_font(18, bold=True)
    label_font = _get_font(12)
    pct_font = _get_font(11, bold=True)

    # Title
    tw, th = _get_text_size(draw, title, title_font)
    draw.text(((width - tw) // 2, 15), title, fill=(0x1B, 0x2A, 0x4A), font=title_font)

    if not values or sum(values) == 0:
        return img

    total = sum(values)
    center_x = width // 2 - 60
    center_y = height // 2 + 15
    radius = min(width, height) // 2 - 80

    # Draw pie slices
    start_angle = -90  # Start from top
    for i, (label, val) in enumerate(zip(labels, values)):
        sweep = (val / total) * 360
        color = CHART_COLORS[i % len(CHART_COLORS)]
        bbox = [center_x - radius, center_y - radius, center_x + radius, center_y + radius]
        draw.pieslice(bbox, start_angle, start_angle + sweep, fill=color, outline=(255, 255, 255), width=2)
        start_angle += sweep

    # Legend on right side
    legend_x = center_x + radius + 40
    legend_y = center_y - (len(labels) * 25) // 2
    for i, (label, val) in enumerate(zip(labels, values)):
        color = CHART_COLORS[i % len(CHART_COLORS)]
        y = legend_y + i * 28
        draw.rectangle([legend_x, y, legend_x + 16, y + 16], fill=color)
        pct = (val / total) * 100
        legend_text = f"{label} ({pct:.1f}%)"
        draw.text((legend_x + 24, y), legend_text, fill=(0x2D, 0x2D, 0x2D), font=label_font)

    return img


def _parse_chart_block(chart_text: str) -> Optional[Dict[str, Any]]:
    """
    Parse a chart definition block. Expected format:

    type: bar|line|pie
    title: Chart Title
    data:
      Label1: 100
      Label2: 200
      Label3: 150
    """
    result = {"type": "bar", "title": "Chart", "labels": [], "values": []}

    lines = chart_text.strip().split("\n")
    in_data = False

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        if stripped.lower().startswith("type:"):
            result["type"] = stripped.split(":", 1)[1].strip().lower()
            in_data = False
        elif stripped.lower().startswith("title:"):
            result["title"] = stripped.split(":", 1)[1].strip()
            in_data = False
        elif stripped.lower() == "data:":
            in_data = True
        elif in_data and ":" in stripped:
            # Parse "Label: value"
            parts = stripped.rsplit(":", 1)
            if len(parts) == 2:
                label = parts[0].strip()
                try:
                    value = float(parts[1].strip().replace(",", "").replace("$", "").replace("%", ""))
                    result["labels"].append(label)
                    result["values"].append(value)
                except ValueError:
                    continue

    if not result["labels"]:
        return None
    return result


def _add_chart_to_doc(doc: Document, chart_data: Dict[str, Any]) -> None:
    """Render a chart and add it as an image to the document."""
    if not PIL_AVAILABLE:
        # Fallback: add chart data as a table
        para = doc.add_paragraph()
        run = para.add_run(f"[Chart: {chart_data.get('title', 'Chart')} — install Pillow for visual charts]")
        run.font.italic = True
        run.font.color.rgb = COLORS["light_gray"]
        return

    chart_type = chart_data.get("type", "bar")
    title = chart_data.get("title", "Chart")
    labels = chart_data.get("labels", [])
    values = chart_data.get("values", [])

    try:
        if chart_type == "line":
            chart_img = _render_line_chart(title, labels, values)
        elif chart_type == "pie":
            chart_img = _render_pie_chart(title, labels, values)
        else:  # default: bar
            chart_img = _render_bar_chart(title, labels, values)

        # Save to temp file and add to doc
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            chart_img.save(tmp, format="PNG", quality=95)
            tmp_path = tmp.name

        # Add centered image
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(tmp_path, width=Inches(5.5))

        # Cleanup temp file
        try:
            os.remove(tmp_path)
        except OSError:
            pass

        # Small spacing after chart
        doc.add_paragraph("")

    except Exception as e:
        logger.warning(f"Failed to render chart: {e}")
        para = doc.add_paragraph()
        run = para.add_run(f"[Chart rendering failed: {str(e)}]")
        run.font.italic = True
        run.font.color.rgb = COLORS["light_gray"]


def _render_section_content(doc: Document, content: str) -> None:
    """
    Parse section content with markdown-like syntax and render into the Word document.
    Supports: sub-headings, bold, italic, tables, bullets, numbered lists,
    blockquotes, horizontal rules, and ```chart blocks.
    """
    lines = content.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # ```chart block — collect until closing ```
        if line.startswith("```chart"):
            i += 1  # skip the opening ```chart line
            chart_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                chart_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # skip closing ```
            chart_text = "\n".join(chart_lines)
            chart_data = _parse_chart_block(chart_text)
            if chart_data:
                _add_chart_to_doc(doc, chart_data)
            continue

        # --- Horizontal rule ---
        if re.match(r'^(---+|___+|\*\*\*+)$', line):
            _add_horizontal_rule(doc)
            i += 1
            continue

        # ### Sub-sub-heading (H3)
        if line.startswith("### "):
            heading_text = line[4:].strip()
            doc.add_heading(heading_text, level=3)
            i += 1
            continue

        # ## Sub-heading (H2)
        if line.startswith("## "):
            heading_text = line[3:].strip()
            doc.add_heading(heading_text, level=2)
            i += 1
            continue

        # > Blockquote
        if line.startswith("> "):
            quote_text = line[2:].strip()
            _add_blockquote(doc, quote_text)
            i += 1
            continue

        # | Table row — collect consecutive table lines
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            _parse_table_block(doc, table_lines)
            continue

        # - Bullet point or * Bullet point
        if line.startswith("- ") or line.startswith("* "):
            bullet_text = line[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            _parse_inline_markdown(p, bullet_text)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(1)
            i += 1
            continue

        # 1. Numbered list
        if re.match(r'^\d+\.\s+', line):
            text = re.sub(r'^\d+\.\s+', '', line)
            p = doc.add_paragraph(style="List Number")
            _parse_inline_markdown(p, text)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(1)
            i += 1
            continue

        # Regular paragraph with inline markdown
        p = doc.add_paragraph()
        _parse_inline_markdown(p, line)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        i += 1


# ---------------------------------------------------------------------------
# Core logic
# ---------------------------------------------------------------------------


async def create_word_db(
    input_data: CreateWordInput,
    sql_db: Session,
    user_id: str,
    conversation_id: str,
) -> str:
    """Generate a professional .docx file from structured section data and persist a DB record."""
    try:
        # 1. Validate user / org
        user = sql_db.query(User).filter(User.id == user_id).first()
        if not user or not user.org_id:
            return "❌ User or Organization not found."

        # 2. Validate section count
        if len(input_data.sections) > MAX_SECTIONS:
            return f"❌ Too many sections ({len(input_data.sections)}). Maximum is {MAX_SECTIONS}."

        # 3. Prepare output path
        sanitized_title = re.sub(r'[^\w\s-]', '', input_data.title).strip().replace(' ', '_')
        if not sanitized_title:
            sanitized_title = "document"

        unique_suffix = str(uuid.uuid4())[:8]
        filename = f"{sanitized_title}_{unique_suffix}.docx"

        storage_rel = f"storage/word_docs/{filename}"
        os.makedirs("storage/word_docs", exist_ok=True)

        # 4. Build Word document
        doc = Document()

        # -- Setup professional styles
        _setup_styles(doc)

        # -- Page margins
        for section_obj in doc.sections:
            section_obj.top_margin = Inches(0.8)
            section_obj.bottom_margin = Inches(0.8)
            section_obj.left_margin = Inches(1.0)
            section_obj.right_margin = Inches(1.0)

        # -- Cover page
        _add_cover_page(doc, input_data.title)

        # -- Table of Contents (only if 3+ sections)
        if len(input_data.sections) >= 3:
            _add_table_of_contents(doc, input_data.sections)

        # -- Header & Footer (applied to all sections)
        _add_header_footer(doc, input_data.title)

        # -- Render each section
        for idx, sec in enumerate(input_data.sections):
            # Section heading
            doc.add_heading(sec.title, level=1)

            # Render content with full markdown support
            _render_section_content(doc, sec.content)

            # Add subtle divider between sections (not after last)
            if idx < len(input_data.sections) - 1:
                _add_section_divider(doc)

        # 5. Save
        doc.save(storage_rel)

        # 6. Check file size
        file_size = os.path.getsize(storage_rel)
        if file_size > MAX_FILE_SIZE_BYTES:
            os.remove(storage_rel)
            return (
                f"❌ Generated file is {file_size / (1024 * 1024):.1f} MB which "
                f"exceeds the 10 MB limit. Please reduce the content."
            )

        # 7. Build public URL
        base = input_data.base_url.rstrip("/")
        file_url = f"{base}/{storage_rel}"

        # 8. Persist DB record
        record = GeneratedWord(
            conversation_id=conversation_id,
            org_id=user.org_id,
            title=input_data.title,
            file_path=storage_rel,
            file_url=file_url,
        )
        sql_db.add(record)
        sql_db.flush()

        for i, sec_in in enumerate(input_data.sections):
            section_record = WordSection(
                generated_word_id=record.id,
                title=sec_in.title,
                content=sec_in.content,
                section_type=sec_in.section_type,
                order=i + 1,
            )
            sql_db.add(section_record)

        sql_db.commit()
        sql_db.refresh(record)

        section_names = ", ".join(f"'{s.title}'" for s in input_data.sections)
        return (
            f"✅ Word document **'{record.title}'** created successfully!\n"
            f"📄 Sections: {section_names}\n"
            f"📥 **Download link:** {file_url}"
        )

    except Exception as e:
        sql_db.rollback()
        logger.error("Error creating Word document: %s", e, exc_info=True)
        return f"❌ Error creating Word document: {str(e)}"


def get_word_link_db(sql_db: Session, conversation_id: str) -> str:
    """Return the download link for the latest Word document in this conversation."""
    try:
        record = (
            sql_db.query(GeneratedWord)
            .filter(GeneratedWord.conversation_id == conversation_id)
            .order_by(GeneratedWord.created_at.desc())
            .first()
        )
        if not record:
            return "📋 No Word document found for this conversation. Use 'create_word_doc' to generate one."
        if not record.file_url:
            return "📋 Word document record found but the file has not been generated yet."
        return (
            f"📥 **Word Document:** '{record.title}'\n"
            f"**Download:** {record.file_url}"
        )
    except Exception as e:
        return f"❌ Error retrieving Word document link: {str(e)}"


# ---------------------------------------------------------------------------
# StructuredTool integration (LangChain)
# ---------------------------------------------------------------------------


def word_generation_tool(
    sql_db: Session,
    user_id: str,
    conversation_id: Optional[str] = None,
    base_url: str = "http://localhost:8000",
) -> List[StructuredTool]:
    """
    Returns LangChain StructuredTools for Word document generation.
    """
    if not conversation_id:
        logger.warning(
            "word_generation_tool called without conversation_id. "
            "Word tools will not be available."
        )
        return []

    # ------------------------------------------------------------------
    # create_word_doc
    # ------------------------------------------------------------------

    async def create_word_doc(
        title: str,
        sections: List[Dict[str, Any]],
    ) -> str:
        """
        Generate a professional Word (.docx) document from structured section data
        and return a download link.

        The document will include a cover page, table of contents, professional
        formatting, headers and footers with page numbers.

        CONTENT FORMATTING — the 'content' field supports markdown-like syntax:
          - **bold text** → bold
          - *italic text* → italic
          - ## Sub-Heading → Heading Level 2
          - ### Sub-Sub-Heading → Heading Level 3
          - - item → bullet point
          - 1. item → numbered list
          - > text → blockquote with left border
          - --- → horizontal rule / section divider
          - | Col1 | Col2 | → table (use | separators, add |---|---| for header row)
          - ```chart ... ``` → embedded chart (bar, line, or pie)

        CHART FORMAT (inside content):
          ```chart
          type: bar
          title: Revenue by Quarter
          data:
            Q1: 180
            Q2: 210
            Q3: 245
            Q4: 285
          ```
          Supported types: bar, line, pie

        SECTIONS FORMAT — each item must have:
          - "title"   : section heading (str)
          - "content" : section body text (str, supports above markdown formatting)

        EXAMPLE:
          sections = [
            {
              "title": "Executive Summary",
              "content": "**Revenue grew 15% year-over-year**, driven by...\\n\\n## Key Metrics\\n\\n| Metric | Q3 2024 | Q3 2025 |\\n|---|---|---|\\n| Revenue | $2.1B | $2.4B |\\n| Margin | 18% | 22% |\\n\\nThis represents a *significant improvement*..."
            },
            {
              "title": "Market Analysis",
              "content": "The North American freight rail market...\\n\\n### Regional Breakdown\\n\\n- **Northeast corridor**: 35% market share\\n- **Midwest**: 28% market share\\n\\n> Industry experts predict a 12% CAGR through 2028."
            }
          ]

        LIMITS:
          - Max 50 sections
          - Max 10 MB file size
        """
        parsed_sections = [
            WordSectionInput(**s) if isinstance(s, dict) else s for s in sections
        ]
        input_obj = CreateWordInput(
            title=title,
            sections=parsed_sections,
            base_url=base_url,
        )
        return await create_word_db(input_obj, sql_db, user_id, conversation_id)

    # ------------------------------------------------------------------
    # get_word_link
    # ------------------------------------------------------------------

    def get_word_link() -> str:
        """
        Return the download link for the most recently created Word document
        in this conversation.
        """
        return get_word_link_db(sql_db, conversation_id)

    # ------------------------------------------------------------------
    # Register as StructuredTools
    # ------------------------------------------------------------------

    class _CreateInput(BaseModel):
        title: str = Field(description="Title / label for the Word document.")
        sections: List[Dict[str, Any]] = Field(
            description=(
                "List of section objects. Each must have 'title' (str) and "
                "'content' (str, supports markdown: **bold**, *italic*, ## headings, "
                "tables, bullets, blockquotes, and ```chart blocks for bar/line/pie charts). "
                "Optional 'section_type' (str). "
                "Max 50 sections. Max 10 MB file size."
            )
        )

    return [
        StructuredTool.from_function(
            coroutine=create_word_doc,
            name="create_word_doc",
            description=(
                "Generate a professional Word (.docx) document with cover page, table of contents, "
                "charts, and rich formatting. Pass 'title' (document name) and 'sections' (list of "
                "{title, content} objects). Content supports markdown: **bold**, *italic*, "
                "## sub-headings, | tables |, - bullets, > blockquotes, --- rules, and "
                "```chart blocks (bar, line, pie). "
                "Returns a download link on success. Limits: max 50 sections, 10 MB."
            ),
            args_schema=_CreateInput,
        ),
        StructuredTool.from_function(
            func=get_word_link,
            name="get_word_link",
            description=(
                "Retrieve the download link for the most recent Word document "
                "created in this conversation."
            ),
            args_schema=None,
        ),
    ]
